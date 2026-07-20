# Author positioning content ONCE as structured blocks -> emit .md, .docx, .html
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

RED="C8102E"; CH="1B1B1B"

def blocks_stanford():
    return [
     ("kicker","COMPETITIVE POSITIONING RESEARCH"),
     ("h1","Stanford HAI — How a Human-Centered AI Institution Positions Itself"),
     ("meta","NexDyne Consulting Group · Brand & Positioning · July 2026"),
     ("p","We studied the Stanford Institute for Human-Centered Artificial Intelligence (HAI) because it is the most recognizable brand built entirely around the phrase “human-centered AI.” Understanding exactly how they own that idea — and where an institution’s voice differs from a corporation’s — sharpens how NexDyne should sound every day."),
     ("h2","What Stanford HAI is"),
     ("p","A university research institute (founded 2019, co-led by Fei-Fei Li and John Etchemendy). It advances AI through research, education, and policy. It is not selling a service or a transformation — it is producing knowledge, shaping policy, and convening the field. That distinction drives everything about how it speaks."),
     ("h2","Their exact positioning language"),
     ("quote","Mission: “Advance AI research, education, policy, and practice to improve the human condition.”"),
     ("quote","Vision: AI “guided by its human impact, inspired by human intelligence, and designed to augment, not replace, people.”"),
     ("p","Three stated research pillars: Intelligence (machine understanding of human language and emotion), Augment Human Capabilities (enhance human work rather than replace it), and Human Impact (how AI interacts with institutions and society)."),
     ("h2","How they talk — the voice"),
     ("bullet","They lead with a QUESTION. Their social cover reads: “How is AI intersecting with our world?” A question signals inquiry and humility — exactly right for a research body."),
     ("bullet","They frame protectively: “augment, NOT replace.” The reassurance is that AI will not displace people."),
     ("bullet","They speak in the collective and the aspirational: “improve the human condition,” “serve the collective needs of humanity.”"),
     ("bullet","They name the gap and position as guide: “a widening gap between what AI can do and how prepared we are to manage it.”"),
     ("h2","What we take from them"),
     ("p","Three things are worth borrowing, and one thing we must deliberately NOT copy."),
     ("bullet","BORROW — the discipline of one owned idea. They own “human-centered.” We should own “human intelligence.”"),
     ("bullet","BORROW — restraint and confidence in the visual voice (calm, editorial, one line of type over a strong image)."),
     ("bullet","BORROW — naming the tension in the market (the gap between AI’s capability and organizations’ readiness) — that tension is our entire reason to exist."),
     ("bullet","DO NOT COPY — the QUESTION format and the defensive “not replace” framing. An institution can afford to wonder and reassure. A corporation that clients pay to deliver outcomes must ASSERT and ADVANCE. We answer the question in the client’s business; we don’t pose it on our banner."),
     ("h2","The strategic gap we exploit"),
     ("p","Stanford augments to PROTECT (keep humans in charge; don’t let AI replace them). NexDyne enhances to ADVANCE (human intelligence plus AI, governed well, becomes a competitive engine that ships innovation). That is the difference between reassurance and advantage — and it is the whole of our daily message."),
     ("rule",""),
     ("small","Sources: hai.stanford.edu/about, hai.stanford.edu, Stanford Report (news.stanford.edu). Public materials, retrieved July 2026."),
    ]

def blocks_nexdyne():
    return [
     ("kicker","BRAND POSITIONING — THE DAILY LEAD"),
     ("h1","NexDyne’s Proposition: Enhancing Human Intelligence to Drive Innovation in the Age of AI"),
     ("meta","NexDyne Consulting Group · Brand & Positioning · July 2026"),
     ("p","This is what we lead with every day — in a post, a pitch, a cover image, a first meeting. It is the answer to “what does NexDyne do?” before any capability, industry, or framework is named."),
     ("h2","The one line"),
     ("quote","We enhance human intelligence to drive innovation in the age of artificial intelligence."),
     ("p","Read it closely, because every word is load-bearing:"),
     ("bullet","“Human intelligence” — our subject and our brand. Not AI. The human is the protagonist; AI is the amplifier. This is what makes us different from every AI vendor."),
     ("bullet","“Enhance” — active and additive. We make people more capable; we do not caretake or protect them from technology."),
     ("bullet","“Drive innovation” — the outcome a business pays for. Not “won’t be replaced,” not “safe.” Advantage."),
     ("bullet","“In the age of artificial intelligence” — names the moment without making AI the hero. AI is the context; human intelligence is the answer."),
     ("h2","Where HIG fits (and where it does not)"),
     ("p","HIG — Human Intelligence, Governed — is our operating discipline: the framework that makes the enhancement safe, accountable, and scalable. It is the HOW, not the WHY. We were leading with the framework; that is like a car company leading with its safety cage instead of the drive. Correct order:"),
     ("bullet","LEAD (the why) — We enhance human intelligence to drive innovation."),
     ("bullet","SUPPORT (the how) — We do it through HIG, so the results are governed, defensible, and built to scale."),
     ("p","HIG earns trust once someone is interested. The proposition earns the interest in the first place. Never open with the governance framework again."),
     ("h2","How we differ — the three foils"),
     ("h3","vs. institutions (Stanford HAI and peers)"),
     ("p","They study and guide; they ask “how is AI intersecting with our world?” and frame AI as “augment, not replace.” We are hired to deliver, so we assert and advance: we don’t protect your people from AI — we make your people, amplified by AI and governed by us, outperform. We answer their question inside your business."),
     ("h3","vs. AI vendors and tool-sellers"),
     ("p","They lead with the model, the platform, the automation — the machine is the hero and the human is a cost to remove. We invert it: the human is the hero and AI is the amplifier. That is why our results are defensible and theirs are fragile — judgment, not just software, is what compounds."),
     ("h3","vs. traditional consultancies"),
     ("p","They sell decks and frameworks and leave. We install a capability — human intelligence, governed and scaled — that keeps producing after we go. “From PowerPoint to production” is not a slogan; it is the difference between advice and an engine."),
     ("h2","Message pillars (use these to generate content)"),
     ("bullet","The human is the protagonist. In every story, a person makes a better decision because AI — governed — gave them reach they didn’t have."),
     ("bullet","Enhancement over replacement. We never sell headcount reduction as the point. We sell capability multiplied."),
     ("bullet","Innovation is the outcome. Governance is the enabler. Speed with control, not speed or control."),
     ("bullet","Defensible by design. Every claim, every result, survives the question “says who?” — because judgment and governance sit underneath it."),
     ("h2","Say this / not that"),
     ("bullet","SAY: “We enhance human intelligence to drive innovation.”  NOT: “We do AI governance.”"),
     ("bullet","SAY: “Your people, amplified and governed, outperform.”  NOT: “AI won’t replace your people.”"),
     ("bullet","SAY: “From pilot to production, safely and at speed.”  NOT: “We help you explore AI.”"),
     ("bullet","SAY a statement.  NOT a question. We answer; we don’t wonder."),
     ("h2","Ready-to-use lines (by surface)"),
     ("bullet","Cover / banner: “Using human intelligence to drive innovations.”"),
     ("bullet","Elevator: “NexDyne enhances human intelligence to drive innovation in the age of AI — governed, so it scales.”"),
     ("bullet","LinkedIn bio line: “We make your people, amplified by AI and governed by discipline, outperform.”"),
     ("bullet","One-word essence: Amplify."),
     ("rule",""),
     ("small","Positioning derived from competitive analysis of Stanford HAI (see companion research doc) and NexDyne’s HIG doctrine. NexDyne Consulting Group, July 2026."),
    ]

def emit_md(blocks, path):
    L=[]
    for t,x in blocks:
        if t=="kicker": L.append(f"**{x}**\n")
        elif t=="h1": L.append(f"# {x}\n")
        elif t=="h2": L.append(f"## {x}\n")
        elif t=="h3": L.append(f"### {x}\n")
        elif t=="meta": L.append(f"*{x}*\n")
        elif t=="p": L.append(f"{x}\n")
        elif t=="bullet": L.append(f"- {x}")
        elif t=="quote": L.append(f"> {x}\n")
        elif t=="small": L.append(f"\n<sub>{x}</sub>\n")
        elif t=="rule": L.append("\n---\n")
    open(path,"w").write("\n".join(L)+"\n")

def emit_docx(blocks, path, title):
    d=Document()
    st=d.styles["Normal"]; st.font.name="Calibri"; st.font.size=Pt(11)
    for t,x in blocks:
        if t=="kicker":
            p=d.add_paragraph(); r=p.add_run(x); r.bold=True; r.font.size=Pt(9); r.font.color.rgb=RGBColor.from_string(RED)
        elif t=="h1":
            p=d.add_paragraph(); r=p.add_run(x); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=RGBColor.from_string(CH)
        elif t=="h2":
            p=d.add_paragraph(); p.paragraph_format.space_before=Pt(12); r=p.add_run(x); r.bold=True; r.font.size=Pt(14); r.font.color.rgb=RGBColor.from_string(CH)
        elif t=="h3":
            p=d.add_paragraph(); r=p.add_run(x); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=RGBColor.from_string(RED)
        elif t=="meta":
            p=d.add_paragraph(); r=p.add_run(x); r.italic=True; r.font.size=Pt(10); r.font.color.rgb=RGBColor.from_string("777777")
        elif t=="p":
            d.add_paragraph(x)
        elif t=="bullet":
            d.add_paragraph(x, style="List Bullet")
        elif t=="quote":
            p=d.add_paragraph(); p.paragraph_format.left_indent=Inches(0.3); r=p.add_run(x); r.italic=True; r.bold=True; r.font.size=Pt(12); r.font.color.rgb=RGBColor.from_string(CH)
        elif t=="small":
            p=d.add_paragraph(); r=p.add_run(x); r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string("999999")
        elif t=="rule":
            d.add_paragraph("—"*30)
    d.save(path)

def emit_html(blocks, path, title):
    css="""<style>
    @page{margin:56px 64px}
    *{box-sizing:border-box} body{font-family:-apple-system,'Helvetica Neue',Arial,sans-serif;color:#1b1b1b;line-height:1.55;max-width:760px;margin:0 auto;padding:40px}
    .kicker{font-size:11px;font-weight:700;letter-spacing:.2em;color:#C8102E;text-transform:uppercase;margin-bottom:6px}
    h1{font-size:30px;line-height:1.15;letter-spacing:-0.02em;margin:0 0 6px}
    h2{font-size:19px;margin:30px 0 8px;letter-spacing:-0.01em}
    h3{font-size:15px;color:#C8102E;margin:20px 0 6px}
    .meta{color:#888;font-style:italic;font-size:13px;margin-bottom:8px}
    p{margin:9px 0}
    ul{margin:9px 0;padding-left:22px} li{margin:5px 0}
    blockquote{border-left:3px solid #C8102E;margin:14px 0;padding:6px 0 6px 18px;font-weight:700;font-size:17px;font-style:italic}
    hr{border:none;border-top:1px solid #ddd;margin:26px 0}
    .small{font-size:11px;color:#999;margin-top:8px}
    </style>"""
    H=[f"<!doctype html><html><head><meta charset='utf-8'><title>{title}</title>{css}</head><body>"]
    inlist=False
    for t,x in blocks:
        if t=="bullet":
            if not inlist: H.append("<ul>"); inlist=True
            H.append(f"<li>{x}</li>"); continue
        if inlist: H.append("</ul>"); inlist=False
        if t=="kicker": H.append(f"<div class='kicker'>{x}</div>")
        elif t=="h1": H.append(f"<h1>{x}</h1>")
        elif t=="h2": H.append(f"<h2>{x}</h2>")
        elif t=="h3": H.append(f"<h3>{x}</h3>")
        elif t=="meta": H.append(f"<div class='meta'>{x}</div>")
        elif t=="p": H.append(f"<p>{x}</p>")
        elif t=="quote": H.append(f"<blockquote>{x}</blockquote>")
        elif t=="small": H.append(f"<div class='small'>{x}</div>")
        elif t=="rule": H.append("<hr>")
    if inlist: H.append("</ul>")
    H.append("</body></html>")
    open(path,"w").write("".join(H))

for name,blocks,title in [
  ("stanford-hai-research", blocks_stanford(), "Stanford HAI — Positioning Research"),
  ("nexdyne-positioning", blocks_nexdyne(), "NexDyne Positioning — The Daily Lead"),
]:
    emit_md(blocks, f"positioning/{name}.md")
    emit_docx(blocks, f"positioning/{name}.docx", title)
    emit_html(blocks, f"positioning/{name}.html", title)
    print("emitted", name, "(.md .docx .html)")
